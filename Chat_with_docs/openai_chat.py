#!/usr/bin/env python3
"""
Simple OpenAI Document Chat Script
A command-line tool to chat with documents using OpenAI's GPT-4.
"""

import os
import time
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class OpenAIDocumentChat:
    def __init__(self, api_key=None, extraction_model="gpt-4.1-mini-2025-04-14", response_model="o3-2025-04-16", fallback_model="gpt-4o"):
        """Initialize the OpenAI client with configurable models."""
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("No API key provided. Set OPENAI_API_KEY environment variable or pass api_key parameter.")
        
        self.client = OpenAI(api_key=self.api_key)
        self.documents = []
        self.chat_history = []
        self.uploaded_files = []
        
        # Model configuration
        self.extraction_model = extraction_model  # Model for extracting relevant context
        self.response_model = response_model      # Model for generating final response
        self.fallback_model = fallback_model      # Model for fallback when two-stage fails
        
        print(f"🤖 Configured models:")
        print(f"   📝 Context Extraction: {self.extraction_model}")
        print(f"   🎯 Response Generation: {self.response_model}")
        print(f"   🔄 Fallback: {self.fallback_model}")
    
    def extract_text_from_docx(self, file_path):
        """Extract text content from DOCX file."""
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return '\n\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {e}")
    
    def load_document(self, file_path):
        """Load a document and prepare it for OpenAI."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_name = os.path.basename(file_path)
        file_extension = file_name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            # Upload PDF file to OpenAI
            try:
                with open(file_path, 'rb') as f:
                    uploaded_file = self.client.files.create(
                        file=f,
                        purpose="user_data"
                    )
                self.uploaded_files.append(uploaded_file)
                
                document = {
                    'type': 'file',
                    'name': file_name,
                    'content': uploaded_file
                }
            except Exception as e:
                raise Exception(f"Error uploading PDF: {e}")
        
        elif file_extension == 'docx':
            # DOCX files need text extraction
            extracted_text = self.extract_text_from_docx(file_path)
            document = {
                'type': 'text',
                'name': file_name,
                'content': f"--- Content from {file_name} ---\n{extracted_text}"
            }
        
        elif file_extension in ['txt', 'md']:
            # Text files are sent as text content
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            document = {
                'type': 'text',
                'name': file_name,
                'content': f"--- Content from {file_name} ---\n{text_content}"
            }
        
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
        
        self.documents.append(document)
        print(f"✅ Loaded document: {file_name}")
        return document
    
    def create_openai_content(self, user_message):
        """Create content array for OpenAI API with documents and text."""
        content_parts = []
        
        # Add uploaded files (PDFs) first
        for doc in self.documents:
            if doc['type'] == 'file':
                content_parts.append({
                    "type": "file",
                    "file": {
                        "file_id": doc['content'].id
                    }
                })
        
        # Combine user message with text content (from DOCX, TXT, MD files)
        text_parts = [user_message]
        for doc in self.documents:
            if doc['type'] == 'text':
                text_parts.append(doc['content'])
        
        # Add combined text content
        content_parts.append({
            "type": "text",
            "text": "\n\n".join(text_parts)
        })
        
        return content_parts
    
    def extract_relevant_context(self, user_query):
        """Extract relevant context from documents using extraction model."""
        if not self.documents:
            return ""
        
        try:
            # Create content parts for extraction model (similar to create_openai_content)
            content_parts = []
            
            # Add uploaded files (PDFs) first
            for doc in self.documents:
                if doc['type'] == 'file':
                    content_parts.append({
                        "type": "file",
                        "file": {
                            "file_id": doc['content'].id
                        }
                    })
            
            # Combine user query with text content and extraction instructions
            text_parts = [f"""You are a document analysis assistant. Your task is to extract only the most relevant sections from the provided documents that would help answer the user's query.

User Query: {user_query}

Instructions:
1. Analyze the user's query carefully
2. Extract only the relevant text sections, paragraphs, or data points from the documents
3. Include enough context but be concise - focus on relevance
4. If multiple documents contain relevant information, include relevant parts from each
5. Maintain the original meaning and context
6. Return only the extracted relevant content, no analysis or summary

Documents to analyze:"""]
            
            # Add text documents
            for doc in self.documents:
                if doc['type'] == 'text':
                    text_parts.append(doc['content'])
            
            # Add combined text content
            content_parts.append({
                "type": "text",
                "text": "\n\n".join(text_parts)
            })

            # Call extraction model for context extraction
            response = self.client.chat.completions.create(
                model=self.extraction_model,
                messages=[{
                    "role": "user", 
                    "content": content_parts
                }],
                max_tokens=4000
            )
            
            extracted_context = response.choices[0].message.content
            return extracted_context.strip()
            
        except Exception as e:
            print(f"⚠️ Error extracting context: {e}")
            # Fallback to original approach if extraction fails
            text_parts = []
            for doc in self.documents:
                if doc['type'] == 'text':
                    text_parts.append(doc['content'])
            return "\n\n".join(text_parts)
    
    def generate_final_response(self, user_query, relevant_context):
        """Generate final response using response model with extracted context."""
        try:
            # Create the final prompt with extracted context
            final_prompt = f"""Based on the relevant context extracted from the documents, please provide a comprehensive and accurate answer to the user's query.

User Query: {user_query}

Relevant Context:
{relevant_context}

Please provide a detailed, well-structured response that directly addresses the user's question using the provided context."""

            # For the second stage, we only send text content (no files)
            # since all relevant information is already extracted in relevant_context
            content_parts = [{
                "type": "text",
                "text": final_prompt
            }]
            
            # Build messages for o3-mini
            messages = []
            
            # Add conversation history (simplified for context)
            for msg in self.chat_history:
                if msg["role"] == "user":
                    messages.append({
                        "role": "user",
                        "content": msg["content"]
                    })
                else:
                    messages.append({
                        "role": "assistant", 
                        "content": msg["content"]
                    })
            
            # Add current message
            messages.append({
                "role": "user",
                "content": content_parts
            })
            
            # Call response model for final response
            stream = self.client.chat.completions.create(
                model=self.response_model,
                messages=messages,
                stream=True
            )
            
            # Stream the response
            full_response = ""
            for chunk in stream:
                if chunk.choices[0].delta.content is not None:
                    chunk_content = chunk.choices[0].delta.content
                    print(chunk_content, end="", flush=True)
                    full_response += chunk_content
            
            return full_response
            
        except Exception as e:
            raise Exception(f"Error generating final response: {e}")
    
    def chat(self, message):
        """Two-stage pipeline: extract context with mini, generate response with o3-mini."""
        pipeline_start_time = time.time()
        
        try:
            print("🔍 Analyzing documents and extracting relevant context...")
            
            # Stage 1: Extract relevant context using extraction model
            stage1_start = time.time()
            relevant_context = self.extract_relevant_context(message)
            stage1_time = time.time() - stage1_start
            
            if relevant_context:
                print(f"✅ Context extracted in {stage1_time:.2f}s. Generating response...")
                print("🤖 OpenAI: ", end="", flush=True)
                
                # Stage 2: Generate final response using response model
                stage2_start = time.time()
                full_response = self.generate_final_response(message, relevant_context)
                stage2_time = time.time() - stage2_start
                
                total_time = time.time() - pipeline_start_time
                
                print(f"\n\n⏱️ Pipeline timing:")
                print(f"   📝 Context extraction ({self.extraction_model}): {stage1_time:.2f}s")
                print(f"   🎯 Response generation ({self.response_model}): {stage2_time:.2f}s")
                print(f"   🔄 Total pipeline time: {total_time:.2f}s")
                
                # Add to chat history
                self.chat_history.append({
                    "role": "user",
                    "content": message
                })
                self.chat_history.append({
                    "role": "assistant",
                    "content": full_response
                })
                
                return full_response
            else:
                raise Exception("No relevant context extracted from documents")
            
        except Exception as e:
            # Fallback to original single-stage approach if two-stage fails
            fallback_time = time.time() - pipeline_start_time
            print(f"\n⚠️ Two-stage pipeline failed in {fallback_time:.2f}s ({e}), falling back to standard approach...")
            return self.chat_fallback(message)
    
    def chat_fallback(self, message):
        """Fallback method using original single-stage approach."""
        fallback_start_time = time.time()
        
        try:
            # Create content for current message (documents + question)
            content_parts = self.create_openai_content(message)
            
            # Build messages array including history
            messages = []
            
            # Add conversation history
            for msg in self.chat_history:
                messages.append({
                    "role": msg["role"],
                    "content": msg["content"]
                })
            
            # Add current user message with documents
            messages.append({
                "role": "user",
                "content": content_parts
            })
            
            # Get streaming response from OpenAI
            stream = self.client.chat.completions.create(
                model=self.fallback_model,
                messages=messages,
                stream=True
            )
            
            # Collect the full response while streaming
            full_response = ""
            print("🤖 OpenAI: ", end="", flush=True)
            for chunk in stream:
                if chunk.choices[0].delta.content is not None:
                    chunk_content = chunk.choices[0].delta.content
                    print(chunk_content, end="", flush=True)
                    full_response += chunk_content
            
            fallback_total_time = time.time() - fallback_start_time
            
            print(f"\n\n⏱️ Fallback timing:")
            print(f"   🔄 Single-stage ({self.fallback_model}): {fallback_total_time:.2f}s")
            
            # Add to chat history
            self.chat_history.append({
                "role": "user",
                "content": message
            })
            self.chat_history.append({
                "role": "assistant",
                "content": full_response
            })
            
            return full_response
            
        except Exception as e:
            fallback_error_time = time.time() - fallback_start_time
            raise Exception(f"Error in fallback chat after {fallback_error_time:.2f}s: {e}")
    
    def list_documents(self):
        """List all loaded documents."""
        if not self.documents:
            print("No documents loaded.")
            return
        
        print("\nLoaded documents:")
        for i, doc in enumerate(self.documents, 1):
            doc_type = "PDF (Uploaded)" if doc['type'] == 'file' else "Text"
            print(f"{i}. {doc['name']} ({doc_type})")
    
    def clear_documents(self):
        """Clear all loaded documents."""
        self.documents = []
        self.uploaded_files = []
        print("All documents cleared.")
    
    def clear_history(self):
        """Clear chat history."""
        self.chat_history = []
        print("Chat history cleared.")
    
    def scan_docs_folder(self, docs_path="docs"):
        """Scan docs folder for supported document types."""
        supported_extensions = ['pdf', 'docx', 'txt', 'md']
        found_files = []
        
        if not os.path.exists(docs_path):
            return found_files
        
        try:
            for filename in os.listdir(docs_path):
                file_path = os.path.join(docs_path, filename)
                if os.path.isfile(file_path):
                    file_extension = filename.split('.')[-1].lower()
                    if file_extension in supported_extensions:
                        found_files.append(file_path)
            
            return sorted(found_files)
        except Exception as e:
            print(f"⚠️  Error scanning docs folder: {e}")
            return found_files
    
    def load_docs_folder(self, docs_path="docs"):
        """Load all supported documents from docs folder."""
        files = self.scan_docs_folder(docs_path)
        loaded_count = 0
        
        for file_path in files:
            try:
                self.load_document(file_path)
                loaded_count += 1
            except Exception as e:
                print(f"⚠️  Could not load {file_path}: {e}")
        
        if loaded_count > 0:
            print(f"📁 Auto-loaded {loaded_count} document(s) from {docs_path} folder")
        
        return loaded_count
    
    def show_model_config(self):
        """Display current model configuration."""
        print("\n🤖 Current Model Configuration:")
        print(f"   📝 Context Extraction: {self.extraction_model}")
        print(f"   🎯 Response Generation: {self.response_model}")
        print(f"   🔄 Fallback: {self.fallback_model}")
        print("\nHow it works:")
        print("   1. Your question is analyzed by the extraction model")
        print("   2. Relevant context is extracted from loaded documents")
        print("   3. The response model generates the final answer using the context")
        print("   4. If any step fails, fallback model handles the complete process")


def main():
    """Main function with command-line interface."""
    print("🤖 OpenAI Document Chat")
    print("=" * 40)
    
    try:
        # Initialize chat client
        chat = OpenAIDocumentChat()
        print("✅ OpenAI client initialized successfully!")
        
        # Auto-load documents from docs folder
        chat.load_docs_folder()
    except Exception as e:
        print(f"❌ Error initializing client: {e}")
        return
    
    
    while True:
        try:
            user_input = input("\n💬 You: ").strip()
            
            if not user_input:
                continue
            
            if user_input.lower() == 'quit':
                print("👋 Goodbye!")
                break
            
            elif user_input.lower().startswith('load '):
                file_path = user_input[5:].strip()
                try:
                    chat.load_document(file_path)
                except Exception as e:
                    print(f"❌ Error loading document: {e}")
            
            elif user_input.lower() == 'list':
                chat.list_documents()
            
            elif user_input.lower() == 'clear':
                chat.clear_documents()
            
            elif user_input.lower() == 'history':
                chat.clear_history()
            
            elif user_input.lower() == 'models':
                chat.show_model_config()
            
            else:
                # Chat with OpenAI
                if not chat.documents:
                    print("⚠️  No documents loaded. Use 'load <file_path>' to load documents first.")
                    continue
                
                print("🤖 OpenAI: ", end="", flush=True)
                try:
                    chat.chat(user_input)
                    print()  # Add newline after streaming is complete
                except Exception as e:
                    print(f"❌ Error: {e}")
        
        except KeyboardInterrupt:
            print("\n👋 Goodbye!")
            break
        except EOFError:
            print("\n👋 Goodbye!")
            break


if __name__ == "__main__":
    main()