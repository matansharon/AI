#!/usr/bin/env python3
"""
Simple Gemini Document Chat Script
A command-line tool to chat with documents using Google's Gemini AI.
"""

import os
from google import genai
import base64
from docx import Document
import io
from dotenv import load_dotenv
import pathlib

# Load environment variables
load_dotenv()

class GeminiDocumentChat:
    def __init__(self, api_key=None):
        """Initialize the Gemini client."""
        self.api_key = api_key or os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not self.api_key:
            raise ValueError("No API key provided. Set GEMINI_API_KEY or GOOGLE_API_KEY environment variable or pass api_key parameter.")
        
        self.client = genai.Client(api_key=self.api_key)
        self.documents = []
        self.chat_history = []
        self.uploaded_files = []
    
    def extract_text_from_docx(self, file_path):
        """Extract text content from DOCX file."""
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return '\n\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {e}")
    
    def load_document(self, file_path):
        """Load a document and prepare it for Gemini."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_name = os.path.basename(file_path)
        file_extension = file_name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            # Upload PDF file to Gemini
            try:
                uploaded_file = self.client.files.upload(file=file_path)
                self.uploaded_files.append(uploaded_file)
                
                document = {
                    'type': 'file',
                    'name': file_name,
                    'content': uploaded_file
                }
            except Exception as e:
                raise Exception(f"Error uploading PDF: {e}")
        
        elif file_extension == 'docx':
            # DOCX files need text extraction
            extracted_text = self.extract_text_from_docx(file_path)
            document = {
                'type': 'text',
                'name': file_name,
                'content': f"--- Content from {file_name} ---\n{extracted_text}"
            }
        
        elif file_extension in ['txt', 'md']:
            # Text files are sent as text content
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            document = {
                'type': 'text',
                'name': file_name,
                'content': f"--- Content from {file_name} ---\n{text_content}"
            }
        
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
        
        self.documents.append(document)
        print(f"✅ Loaded document: {file_name}")
        return document
    
    def create_gemini_content(self, user_message):
        """Create content array for Gemini API with documents and text."""
        content_parts = []
        
        # Add user message first
        content_parts.append(user_message)
        
        # Add uploaded files (PDFs)
        for doc in self.documents:
            if doc['type'] == 'file':
                content_parts.append(doc['content'])
        
        # Add text content (from DOCX, TXT, MD files)
        for doc in self.documents:
            if doc['type'] == 'text':
                content_parts.append(doc['content'])
        
        return content_parts
    
    def chat(self, message):
        """Send a message to Gemini and get a response."""
        try:
            # Create content for current message (documents + question)
            content_parts = self.create_gemini_content(message)
            
            # Get response from Gemini
            response = self.client.models.generate_content(
                model="gemini-2.0-flash",
                contents=content_parts
            )
            
            assistant_response = response.text
            
            # Add to chat history (simplified for Gemini)
            self.chat_history.append({
                "role": "user",
                "content": message
            })
            self.chat_history.append({
                "role": "assistant", 
                "content": assistant_response
            })
            
            return assistant_response
            
        except Exception as e:
            raise Exception(f"Error chatting with Gemini: {e}")
    
    def list_documents(self):
        """List all loaded documents."""
        if not self.documents:
            print("No documents loaded.")
            return
        
        print("\nLoaded documents:")
        for i, doc in enumerate(self.documents, 1):
            doc_type = "PDF (Uploaded)" if doc['type'] == 'file' else "Text"
            print(f"{i}. {doc['name']} ({doc_type})")
    
    def clear_documents(self):
        """Clear all loaded documents."""
        self.documents = []
        self.uploaded_files = []
        print("All documents cleared.")
    
    def clear_history(self):
        """Clear chat history."""
        self.chat_history = []
        print("Chat history cleared.")
    
    def scan_docs_folder(self, docs_path="docs"):
        """Scan docs folder for supported document types."""
        supported_extensions = ['pdf', 'docx', 'txt', 'md']
        found_files = []
        
        if not os.path.exists(docs_path):
            return found_files
        
        try:
            for filename in os.listdir(docs_path):
                file_path = os.path.join(docs_path, filename)
                if os.path.isfile(file_path):
                    file_extension = filename.split('.')[-1].lower()
                    if file_extension in supported_extensions:
                        found_files.append(file_path)
            
            return sorted(found_files)
        except Exception as e:
            print(f"⚠️  Error scanning docs folder: {e}")
            return found_files
    
    def load_docs_folder(self, docs_path="docs"):
        """Load all supported documents from docs folder."""
        files = self.scan_docs_folder(docs_path)
        loaded_count = 0
        
        for file_path in files:
            try:
                self.load_document(file_path)
                loaded_count += 1
            except Exception as e:
                print(f"⚠️  Could not load {file_path}: {e}")
        
        if loaded_count > 0:
            print(f"📁 Auto-loaded {loaded_count} document(s) from {docs_path} folder")
        
        return loaded_count


def main():
    """Main function with command-line interface."""
    print("🤖 Gemini Document Chat")
    print("=" * 40)
    
    try:
        # Initialize chat client
        chat = GeminiDocumentChat()
        print("✅ Gemini client initialized successfully!")
        
        # Auto-load documents from docs folder
        chat.load_docs_folder()
    except Exception as e:
        print(f"❌ Error initializing client: {e}")
        return
    
    print("\nCommands:")
    print("  load <file_path>  - Load a document")
    print("  list             - List loaded documents")
    print("  clear            - Clear all documents")
    print("  history          - Clear chat history")
    print("  quit             - Exit the program")
    print("  Or just type your question to chat with documents")
    print("-" * 40)
    
    while True:
        try:
            user_input = input("\n💬 You: ").strip()
            
            if not user_input:
                continue
            
            if user_input.lower() == 'quit':
                print("👋 Goodbye!")
                break
            
            elif user_input.lower().startswith('load '):
                file_path = user_input[5:].strip()
                try:
                    chat.load_document(file_path)
                except Exception as e:
                    print(f"❌ Error loading document: {e}")
            
            elif user_input.lower() == 'list':
                chat.list_documents()
            
            elif user_input.lower() == 'clear':
                chat.clear_documents()
            
            elif user_input.lower() == 'history':
                chat.clear_history()
            
            else:
                # Chat with Gemini
                if not chat.documents:
                    print("⚠️  No documents loaded. Use 'load <file_path>' to load documents first.")
                    continue
                
                print("🤖 Gemini: ", end="", flush=True)
                try:
                    response = chat.chat(user_input)
                    print(response)
                except Exception as e:
                    print(f"❌ Error: {e}")
        
        except KeyboardInterrupt:
            print("\n👋 Goodbye!")
            break
        except EOFError:
            print("\n👋 Goodbye!")
            break


if __name__ == "__main__":
    main()