import streamlit as st
import os
import anthropic
import tempfile
from dotenv import load_dotenv
from docx import Document
import io
import base64

# Load environment variables
load_dotenv()

st.set_page_config(
    page_title="Chat with Documents - Claude AI",
    page_icon="📄",
    layout="wide"
)

def initialize_client(api_key):
    """Initialize the Claude client with provided API key."""
    try:
        return anthropic.Anthropic(api_key=api_key)
    except Exception as e:
        st.error(f"Failed to initialize client: {e}")
        return None

def extract_text_from_docx(file_bytes):
    """Extract text content from DOCX file bytes."""
    try:
        # Create a Document object from file bytes
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {e}")

def prepare_document_for_claude(file, file_name, file_extension):
    """Prepare a document for Claude based on its type."""
    file_content = file.getvalue()
    
    if file_extension.lower() == 'pdf':
        # PDFs are sent as documents with base64 encoding
        pdf_data = base64.standard_b64encode(file_content).decode("utf-8")
        return {
            'type': 'document',
            'name': file_name,
            'content': {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_data
                }
            }
        }
    elif file_extension.lower() == 'docx':
        # DOCX files need text extraction
        extracted_text = extract_text_from_docx(file_content)
        return {
            'type': 'text',
            'name': file_name,
            'content': f"--- Content from {file_name} ---\n{extracted_text}"
        }
    elif file_extension.lower() in ['txt', 'md']:
        # Text files are sent as text content
        text_content = file_content.decode('utf-8')
        return {
            'type': 'text',
            'name': file_name,
            'content': f"--- Content from {file_name} ---\n{text_content}"
        }
    else:
        raise Exception(f"Unsupported file type: {file_extension}")

def create_claude_content(documents, text_message):
    """Create content array for Claude API with documents and text."""
    content = []
    
    # Add documents first (PDFs)
    for doc in documents:
        if doc['type'] == 'document':
            content.append(doc['content'])
    
    # Add text content (from DOCX, TXT, MD files)
    text_parts = [text_message]
    for doc in documents:
        if doc['type'] == 'text':
            text_parts.append(doc['content'])
    
    # Combine all text content into one text block
    if len(text_parts) > 1:
        content.append({
            "type": "text",
            "text": "\n\n".join(text_parts)
        })
    else:
        content.append({
            "type": "text", 
            "text": text_message
        })
    
    return content

def main():
    """Main function for chat with documents."""
    st.title("📄 Chat with Documents - Claude AI")
    st.markdown("Upload documents and have a conversation with Claude about their content")
    st.markdown("---")
    
    # Sidebar for API key and document upload
    with st.sidebar:
        st.header("🔑 Configuration")
        
        # Try to get API key from environment first
        api_key = os.getenv("ANTHROPIC_API_KEY")
        
        
        
        
            
        client = initialize_client(api_key)
        if not client:
            st.error("❌ Failed to initialize client")
            return
        
        st.success("✅ Client initialized successfully!")
        st.markdown("---")
        
        # Document Upload Section
        st.header("📁 Upload Documents")
        uploaded_files = st.file_uploader(
            "Choose documents to chat about:",
            type=['pdf', 'txt', 'docx', 'md'],
            accept_multiple_files=True,
            help="Upload documents you want to discuss with Claude. PDFs are processed natively (including images/charts), while DOCX files have text extracted. TXT/MD files are sent as text content."
        )
        
        # Document management
        if uploaded_files:
            st.markdown("### 📂 Uploaded Documents:")
            for i, file in enumerate(uploaded_files, 1):
                file_extension = file.name.split('.')[-1].lower()
                if file_extension == 'docx':
                    st.markdown(f"**{i}.** {file.name} 📝 (Text will be extracted)")
                else:
                    st.markdown(f"**{i}.** {file.name} 📄")
            
            if st.button("🚀 Start Chat with Documents", type="primary"):
                with st.spinner("Processing documents..."):
                    try:
                        # Initialize session state for document chat
                        st.session_state.documents = []
                        st.session_state.chat_history = []
                        st.session_state.documents_loaded = True
                        
                        # Process all documents for Claude
                        for file in uploaded_files:
                            file_extension = file.name.split('.')[-1].lower()
                            doc = prepare_document_for_claude(file, file.name, file_extension)
                            st.session_state.documents.append(doc)
                        
                        # Create initial context message
                        initial_message = f"I have uploaded {len(uploaded_files)} documents: {', '.join([f.name for f in uploaded_files])}. Please analyze these documents and let me know you're ready to answer questions about them."
                        
                        # Create content for Claude API
                        content = create_claude_content(st.session_state.documents, initial_message)
                        
                        # Create container for streaming response
                        response_container = st.empty()
                        
                        # Start with processing indicator
                        response_container.markdown("**🤖 Claude:** 📄 _Analyzing documents..._")
                        
                        # Send initial message with streaming
                        with client.messages.stream(
                            model="claude-sonnet-4-20250514",
                            max_tokens=1024,
                            messages=[{
                                "role": "user",
                                "content": content
                            }]
                        ) as stream:
                            full_response = ""
                            for chunk in stream.text_stream:
                                full_response += chunk
                                response_container.markdown(f"**🤖 Claude:** {full_response}▋")
                        
                        # Remove cursor when done
                        response_container.markdown(f"**🤖 Claude:** {full_response}")
                        
                        # Add to chat history
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": full_response
                        })
                        
                        st.success(f"✅ {len(uploaded_files)} documents processed successfully!")
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"Error processing documents: {e}")
    
    # Main Chat Interface
    if 'documents_loaded' in st.session_state and st.session_state.documents_loaded:
        st.header("💬 Chat with Your Documents")
        
        # Display document context
        with st.expander("📋 Loaded Documents", expanded=False):
            for doc in st.session_state.documents:
                if doc['type'] == 'document':
                    st.markdown(f"• **{doc['name']}** (PDF - Native Processing)")
                else:
                    st.markdown(f"• **{doc['name']}** (Text Extracted)")
        
        # Document previews (separate from the main expander)
        text_docs = [doc for doc in st.session_state.documents if doc['type'] == 'text']
        if text_docs:
            for doc in text_docs:
                with st.expander(f"📄 Preview: {doc['name']}", expanded=False):
                    # Extract just the content part after the header
                    content_parts = doc['content'].split('---\n', 1)
                    preview_content = content_parts[1] if len(content_parts) > 1 else doc['content']
                    preview = preview_content[:500] + "..." if len(preview_content) > 500 else preview_content
                    st.text(preview)
        
        # Chat controls
        col1, col2 = st.columns([3, 1])
        with col2:
            if st.button("🗑️ Clear Chat"):
                st.session_state.chat_history = []
                st.rerun()
        
        # Chat input
        user_question = st.text_input(
            "Ask a question about your documents:",
            placeholder="What is the main topic of these documents?",
            key="doc_chat_input"
        )
        
        # Add hint
        st.caption("💡 Tip: Type your question and press Enter or click Send")
        
        if st.button("📤 Send Question", type="primary") and user_question.strip():
            try:
                # Show user message immediately
                st.markdown(f"**👤 You:** {user_question}")
                
                # Create placeholder for streaming response
                response_placeholder = st.empty()
                response_placeholder.markdown("**🤖 Claude:** ✨ _Thinking..._")
                
                # Build messages for Claude API
                messages = []
                
                # Add conversation history to messages
                for msg in st.session_state.chat_history:
                    messages.append({
                        "role": msg["role"],
                        "content": msg["content"]
                    })
                
                # Create content for current message (documents + question)
                content = create_claude_content(st.session_state.documents, user_question)
                
                # Add current user message
                messages.append({
                    "role": "user", 
                    "content": content
                })
                
                # Stream response from Claude
                try:
                    with client.messages.stream(
                        model="claude-sonnet-4-20250514",
                        max_tokens=1024,
                        messages=messages
                    ) as stream:
                        full_response = ""
                        for chunk in stream.text_stream:
                            full_response += chunk
                            # Show streaming response with cursor
                            response_placeholder.markdown(f"**🤖 Claude:** {full_response}▋")
                    
                    # Remove cursor when done
                    response_placeholder.markdown(f"**🤖 Claude:** {full_response}")
                    
                except Exception as stream_error:
                    response_placeholder.markdown(f"**🤖 Claude:** ❌ Error: {stream_error}")
                    full_response = f"Error: {stream_error}"
                
                # Add both user question and assistant response to history
                st.session_state.chat_history.append({
                    "role": "user",
                    "content": user_question
                })
                st.session_state.chat_history.append({
                    "role": "assistant", 
                    "content": full_response
                })
                
                # Clear the placeholder and rerun to show full chat history
                response_placeholder.empty()
                st.rerun()
                    
            except Exception as e:
                st.error(f"Error sending message: {e}")
        
        # Display chat history
        if st.session_state.chat_history:
            st.markdown("### 💬 Conversation")
            for message in st.session_state.chat_history:
                if message["role"] == "user":
                    st.markdown(f"**👤 You:** {message['content']}")
                else:
                    st.markdown(f"**🤖 Claude:** {message['content']}")
                st.markdown("---")
    
    else:
        # Welcome message when no documents are loaded
        st.markdown("""
        ## Welcome! 👋
        
        To get started:
        1. **Upload documents** in the sidebar (PDF, TXT, DOCX, MD)
        2. **Click "Start Chat"** to process your documents with Claude
        3. **Ask questions** about your documents in the chat interface
        
        ### What Claude can do:
        - **PDFs**: Native processing including images, charts, and tables
        - **Text files**: Direct analysis of content
        - **DOCX files**: Text extraction and analysis
        
        ### Example Questions:
        - "What are the main topics discussed in these documents?"
        - "Can you summarize the key points?"
        - "What conclusions can you draw from this content?"
        - "Are there any important dates or numbers mentioned?"
        """)
    
    # Footer
    st.markdown("---")
    st.markdown(
        "Made with ❤️ using [Streamlit](https://streamlit.io) and [Anthropic Claude API](https://anthropic.com)"
    )

if __name__ == "__main__":
    main()