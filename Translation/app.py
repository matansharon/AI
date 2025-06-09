from openai import OpenAI
openai_client = OpenAI()
import os, json, base64
from typing import List, Dict, Any

import unstructured_client
from unstructured_client.models import operations, shared
from docx import Document
from docx.shared import Inches
from io import BytesIO

client = unstructured_client.UnstructuredClient(
    api_key_auth=os.getenv("UNSTRUCTURED_API_KEY")
)




def translate_with_openai(text,src_lang='en', tgt_lang='he'):
    completion = openai_client.chat.completions.create(
        model="gpt-4.1-mini-2025-04-14",
        messages=[
            {
                'role': 'assistant',
                'content': f"""You are a professional translator. Your task is to translate the text from {src_lang} to {tgt_lang}. 
                'You must return the translation without any additional text. 
                'You must not return the original text. 
                the document are an engineering document,specific to the field manufacture engendering on the medical devices industry. 
                """
            },
            {
                "role": "user",
                "content": f"Translate this text to Hebrew: {text} just return the translation without any additional text."
            }
        ]
    )

    return completion.choices[0].message.content



def get_list_of_elements_from_file_with_unstructured_client(file_path:str,src_lang='he', tgt_lang='en'):
    """
    This function uses the Unstructured Client to process a file and return a list of elements.
    :param filename: The path to the input file.
    :param src_lang: Source language for translation.
    :param tgt_lang: Target language for translation.
    :return: List of elements processed from the file.
    """
    

    import unstructured_client
    from unstructured_client.models import operations, shared

    client = unstructured_client.UnstructuredClient(
        api_key_auth=os.getenv("UNSTRUCTURED_API_KEY")
    )

    filename = file_path

    req = operations.PartitionRequest(
        partition_parameters=shared.PartitionParameters(
            files=shared.Files(
                content=open(filename, "rb"),
                file_name=filename,
            ),
            strategy=shared.Strategy.HI_RES,
            
        ),
    )

    try:
        res = client.general.partition(
            request=req
        )
        element_dicts = [element for element in res.elements]
        return element_dicts
        
    except Exception as e:
        print(e)


def process_elements_for_translation(elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Process elements to identify which ones need translation and prepare them.
    :param elements: List of elements from unstructured client
    :return: List of processed elements with translation flags
    """
    text_element_types = ['Text', 'Title', 'Header', 'Footer', 'NarrativeText', 'ListItem', 'UncategorizedText']
    
    processed_elements = []
    
    for element in elements:
        processed_element = dict(element)
        element_type = element.get('type', '')
        
        # Determine if element needs translation
        needs_translation = (
            element_type in text_element_types and 
            element.get('text', '').strip() != ''
        )
        
        processed_element['needs_translation'] = needs_translation
        processed_element['original_text'] = element.get('text', '') if needs_translation else None
        processed_element['translated_text'] = None
        
        processed_elements.append(processed_element)
    
    return processed_elements


def translate_elements_batch(elements: List[Dict[str, Any]], src_lang='he', tgt_lang='en') -> List[Dict[str, Any]]:
    """
    Translate all elements that need translation in batch for efficiency.
    :param elements: List of processed elements
    :param src_lang: Source language
    :param tgt_lang: Target language
    :return: List of elements with translations
    """
    translated_elements = []
    
    for element in elements:
        if element.get('needs_translation', False):
            try:
                original_text = element['original_text']
                translated_text = translate_with_openai(original_text, src_lang, tgt_lang)
                element['translated_text'] = translated_text
                element['text'] = translated_text  # Update the main text field
            except Exception as e:
                print(f"Translation error for element: {e}")
                element['translated_text'] = element['original_text']  # Fallback to original
                element['text'] = element['original_text']
        
        translated_elements.append(element)
    
    return translated_elements


def extract_image_data(element: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract image data from element metadata.
    :param element: Element that might contain image data
    :return: Dictionary with image info (data, format, etc.)
    """
    image_info = {
        'has_image': False,
        'image_data': None,
        'image_format': None,
        'width': None,
        'height': None
    }
    
    if element.get('type') == 'Image':
        metadata = element.get('metadata', {})
        
        # Look for base64 image data in various possible keys
        for key in ['image_base64', 'base64_image', 'image_data', 'base64']:
            if key in metadata and metadata[key]:
                image_info['has_image'] = True
                image_info['image_data'] = metadata[key]
                break
        
        # Extract image format and dimensions if available
        if 'filetype' in metadata:
            image_info['image_format'] = metadata['filetype']
        elif 'format' in metadata:
            image_info['image_format'] = metadata['format']
        
        if 'width' in metadata:
            image_info['width'] = metadata['width']
        if 'height' in metadata:
            image_info['height'] = metadata['height']
    
    return image_info


def build_docx_from_elements(elements: List[Dict[str, Any]], output_path: str) -> bool:
    """
    Build a new DOCX file from the translated elements.
    :param elements: List of translated elements
    :param output_path: Path for the output DOCX file
    :return: True if successful, False otherwise
    """
    try:
        doc = Document()
        
        for element in elements:
            element_type = element.get('type', '')
            text = element.get('text', '')
            
            if element_type == 'Title':
                # Add as heading level 1
                doc.add_heading(text, level=1)
                
            elif element_type == 'Header':
                # Add as heading level 2
                doc.add_heading(text, level=2)
                
            elif element_type in ['Text', 'NarrativeText', 'UncategorizedText']:
                # Add as regular paragraph
                if text.strip():
                    doc.add_paragraph(text)
                    
            elif element_type == 'ListItem':
                # Add as bullet point
                if text.strip():
                    p = doc.add_paragraph(text, style='ListBullet')
                    
            elif element_type == 'Table':
                # Handle table elements
                add_table_to_doc(doc, element)
                
            elif element_type == 'Image':
                # Handle image elements
                add_image_to_doc(doc, element)
                
            elif element_type == 'Footer':
                # Add footer text as italic paragraph
                if text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.italic = True
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error building DOCX: {e}")
        return False


def add_table_to_doc(doc: Document, element: Dict[str, Any]):
    """
    Add a table element to the document.
    :param doc: Document object
    :param element: Table element from unstructured
    """
    try:
        # For now, add table as text - can be enhanced later
        text = element.get('text', '')
        if text.strip():
            p = doc.add_paragraph(text)
            p.style = 'Normal'
    except Exception as e:
        print(f"Error adding table: {e}")


def add_image_to_doc(doc: Document, element: Dict[str, Any]):
    """
    Add an image element to the document.
    :param doc: Document object
    :param element: Image element from unstructured
    """
    try:
        image_info = extract_image_data(element)
        
        if image_info['has_image'] and image_info['image_data']:
            # Decode base64 image data
            image_data = base64.b64decode(image_info['image_data'])
            image_stream = BytesIO(image_data)
            
            # Add image to document
            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Set image size (default to 4 inches width if no dimensions available)
            width = Inches(4)
            if image_info['width']:
                try:
                    width = Inches(min(6, max(1, image_info['width'] / 72)))  # Convert and clamp
                except:
                    width = Inches(4)
            
            run.add_picture(image_stream, width=width)
            
    except Exception as e:
        print(f"Error adding image: {e}")


def translate_document_workflow(input_file_path: str, output_file_path: str, src_lang='he', tgt_lang='en') -> bool:
    """
    Complete workflow to translate a document from input to output.
    :param input_file_path: Path to input DOCX or PDF file
    :param output_file_path: Path for translated DOCX output file
    :param src_lang: Source language code
    :param tgt_lang: Target language code
    :return: True if successful, False otherwise
    """
    try:
        print(f"Starting translation workflow: {input_file_path} -> {output_file_path}")
        
        # Step 1: Extract elements from file using unstructured
        print("Step 1: Extracting elements from file...")
        raw_elements = get_list_of_elements_from_file_with_unstructured_client(input_file_path, src_lang, tgt_lang)
        
        if not raw_elements:
            print("Error: No elements extracted from file")
            return False
        
        print(f"Extracted {len(raw_elements)} elements")
        
        # Step 2: Process elements for translation
        print("Step 2: Processing elements for translation...")
        processed_elements = process_elements_for_translation(raw_elements)
        
        translatable_count = sum(1 for el in processed_elements if el.get('needs_translation', False))
        print(f"Found {translatable_count} elements that need translation")
        
        # Step 3: Translate text elements
        print("Step 3: Translating text elements...")
        translated_elements = translate_elements_batch(processed_elements, src_lang, tgt_lang)
        
        # Step 4: Build output DOCX file
        print("Step 4: Building output DOCX file...")
        success = build_docx_from_elements(translated_elements, output_file_path)
        
        if success:
            print(f"Translation completed successfully! Output saved to: {output_file_path}")
            return True
        else:
            print("Error: Failed to build output DOCX file")
            return False
            
    except Exception as e:
        print(f"Error in translation workflow: {e}")
        return False


# Example usage
if __name__ == "__main__":
    # Example: Translate the Hebrew document to English
    input_file = "input.docx"
    output_file = "translated_document_en.docx"
    
    success = translate_document_workflow(
        input_file_path=input_file,
        output_file_path=output_file,
        src_lang='he',  # Hebrew
        tgt_lang='en'   # English
    )
    
    if success:
        print("Document translation completed successfully!")
    else:
        print("Document translation failed. Check error messages above.")